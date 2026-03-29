#!/usr/bin/env python3
"""
Convert cognitive dismantling books from markdown to docx format.
Creates one docx file per workbook (subfolder), combining all chapters.
Handles markdown formatting: headings, bold, italic, bullet points, numbered lists.
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def parse_inline_formatting(paragraph, text):
    """
    Parse text with inline markdown formatting and add to paragraph.
    Handles: **bold**, *italic*, ***bold+italic***
    """
    # Pattern to match bold, italic, or bold+italic
    pattern = r'(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*)'
    
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
            
        if part.startswith('***') and part.endswith('***'):
            # Bold + italic
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Regular text
            paragraph.add_run(part)


def add_heading(doc, text, level):
    """Add a heading with proper formatting."""
    # Remove markdown heading markers
    clean_text = text.lstrip('#').strip()
    heading = doc.add_heading(clean_text, level=level)
    return heading


def add_paragraph(doc, text):
    """Add a regular paragraph with inline formatting."""
    paragraph = doc.add_paragraph()
    parse_inline_formatting(paragraph, text)
    return paragraph


def add_bullet_list(doc, items):
    """Add a bullet list."""
    for item in items:
        paragraph = doc.add_paragraph(style='List Bullet')
        parse_inline_formatting(paragraph, item)


def add_numbered_list(doc, items):
    """Add a numbered list."""
    for item in items:
        paragraph = doc.add_paragraph(style='List Number')
        parse_inline_formatting(paragraph, item)


def process_markdown_file(doc, md_path):
    """
    Process a single markdown file and add its content to the document.
    """
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Handle headings
        if line.startswith('#'):
            # Count heading level
            level = len(line) - len(line.lstrip('#'))
            level = min(level, 9)  # Word supports up to 9 heading levels
            add_heading(doc, line, level)
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() == '---':
            # Add a paragraph break
            doc.add_paragraph()
            i += 1
            continue
        
        # Handle bullet lists
        if line.strip().startswith('* ') or line.strip().startswith('- '):
            bullet_items = []
            while i < len(lines):
                current = lines[i]
                if current.strip().startswith('* ') or current.strip().startswith('- '):
                    # Remove bullet marker
                    item = current.strip()[2:]
                    bullet_items.append(item)
                    i += 1
                elif current.strip() == '':
                    i += 1
                    break
                else:
                    break
            add_bullet_list(doc, bullet_items)
            continue
        
        # Handle numbered lists
        numbered_match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if numbered_match:
            numbered_items = []
            while i < len(lines):
                current = lines[i]
                match = re.match(r'^(\d+)\.\s+(.+)$', current.strip())
                if match:
                    numbered_items.append(match.group(2))
                    i += 1
                elif current.strip() == '':
                    i += 1
                    break
                else:
                    break
            add_numbered_list(doc, numbered_items)
            continue
        
        # Handle indented bullet points (like in chapter_01.md)
        if line.startswith('    * ') or line.startswith('    - '):
            # This is a sub-bullet, treat as regular bullet for now
            item = line.strip()[2:]
            paragraph = doc.add_paragraph(style='List Bullet 2')
            parse_inline_formatting(paragraph, item)
            i += 1
            continue
        
        # Regular paragraph
        paragraph_text = line
        i += 1
        
        # Continue reading paragraph until empty line or special element
        while i < len(lines):
            next_line = lines[i]
            if (next_line.strip() == '' or 
                next_line.startswith('#') or 
                next_line.strip() == '---' or
                next_line.strip().startswith('* ') or
                next_line.strip().startswith('- ') or
                re.match(r'^(\d+)\.\s+', next_line.strip())):
                break
            paragraph_text += ' ' + next_line
            i += 1
        
        add_paragraph(doc, paragraph_text)


def create_workbook_docx(workbook_path, output_path):
    """
    Create a docx file for a workbook directory.
    Combines all chapter files in order.
    """
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Get all markdown files in the directory
    md_files = []
    for file in os.listdir(workbook_path):
        if file.endswith('.md') and (file.startswith('chapter_') or file == 'finale.md'):
            md_files.append(file)
    
    # Sort files: chapters first (numerically), then finale
    def sort_key(filename):
        if filename == 'finale.md':
            return (1, 999)  # Finale goes last
        match = re.match(r'chapter_(\d+)\.md', filename)
        if match:
            return (0, int(match.group(1)))
        return (0, 0)
    
    md_files.sort(key=sort_key)
    
    # Process each file
    for md_file in md_files:
        md_path = os.path.join(workbook_path, md_file)
        print(f"  Processing: {md_file}")
        process_markdown_file(doc, md_path)
        
        # Add page break between chapters (except after finale)
        if md_file != 'finale.md':
            doc.add_page_break()
    
    # Save the document
    doc.save(output_path)
    print(f"  Created: {output_path}")


def main():
    """Main function to process all workbooks."""
    base_dir = Path('cognitive_dismantling_books')
    
    if not base_dir.exists():
        print(f"Error: Directory '{base_dir}' not found.")
        return
    
    # Get all subdirectories (workbooks)
    workbooks = [d for d in base_dir.iterdir() if d.is_dir()]
    
    print(f"Found {len(workbooks)} workbooks to process:\n")
    
    for workbook in workbooks:
        workbook_name = workbook.name
        print(f"\nProcessing: {workbook_name}")
        
        # Create output filename
        output_filename = f"{workbook_name}.docx"
        output_path = workbook / output_filename
        
        create_workbook_docx(str(workbook), str(output_path))
    
    print("\n\nConversion complete!")


if __name__ == '__main__':
    main()
